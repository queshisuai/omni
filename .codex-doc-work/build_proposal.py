from copy import deepcopy
from pathlib import Path
import hashlib

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REFERENCE = Path(r"C:\Users\Administrator\Desktop\开题报告\广州工商学院本科毕业论文（设计）开题报告 .docx")
OUTPUT = Path(r"C:\Users\Administrator\Desktop\开题报告\202316510149-余凯欣-开题报告-第3稿.docx")
EXPECTED_HASH = "5A412A01F863C53F81AD53CAF49F596613EB7037D40A87E7F7929F76EF748190"


REFERENCES = [
    "中国演出行业协会，灯塔专业版．2024大型营业性演出市场趋势及特点分析[R]．北京：中国演出行业协会，2025．",
    "何锋，罗胜，罗丽娟．微服务架构的一体化性能监控SaaS云设计与实现[J]．计算机应用与软件，2024，41(8)：28-35．",
    "李淑霞，赵泽龙，孙海萍，等．基于Spring Cloud微服务架构的能源互联网营销服务系统设计[J]．信息技术，2025(10)：138-145．",
    "庞长才．基于云原生技术的管理信息系统微服务架构设计与实现[J]．中国新通信，2026，28(3)：16-18，24．",
    "张健．基于Spring Cloud微服务架构的工业软件多层级组件平台设计[J]．自动化与仪器仪表，2026(1)：131-134，139．",
    "张佳佳，谭艳．构建政务信息服务一体化平台的关键技术及实施策略[J]．电子通信与计算机科学，2025，7(7)：22-24．",
    "陈彦成．微服务架构下的Web应用开发及跨域通信解决方案[J]．电子通信与计算机科学，2025，7(8)：212-214．",
    "甘日进．高并发分布式系统的负载均衡机制设计[J]．计算机与自主智能研究进展，2025，3(4)：30．",
    "王佳艺．民航旅客服务系统的架构演变[J]．电子通信与计算机科学，2025，7(10)：184-186．",
    "唐俊，邓东杰．基于微服务架构的气象信息共享平台重构与服务治理[J]．科技创新发展，2026，3(7)．",
    "GUTIERREZ F．Spring Cloud with Spring Boot[A]//Pro Spring Boot 3[M]．Berkeley：Apress，2024：701-799．",
    "GUO C，SHI X，JI M，et al．Design and Implementation of Intelligent Logistics Control Platform Based on Spring Cloud[J]．Procedia Computer Science，2024，247：529-536．",
    "CHEN J，FAN R，SHAO C，et al．Optimisation Strategies for Load Balancing Algorithms Based on Spring Cloud Alibaba[C]//Proceedings of the 2024 3rd Asia Conference on Algorithms，Computing and Machine Learning．New York：ACM，2024：207-211．",
]


def set_run_font(run, chinese="宋体", western="Times New Roman", size=Pt(12), bold=False, superscript=False):
    run.font.name = western
    run.font.size = size
    run.font.bold = bold
    run.font.superscript = superscript
    run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese)
    run._element.rPr.rFonts.set(qn("w:ascii"), western)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), western)


def set_para_format(paragraph, indent=True, line_spacing=1.5, before=0, after=0, keep_with_next=False):
    fmt = paragraph.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = Pt(24) if indent else Pt(0)
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.keep_with_next = keep_with_next
    fmt.keep_together = False
    fmt.widow_control = True


def clear_cell(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    tc.append(OxmlElement("w:p"))


def add_heading(cell, text, level=1):
    p = cell.add_paragraph()
    set_para_format(p, indent=False, line_spacing=1.5, before=3 if level == 1 else 0, after=0, keep_with_next=True)
    run = p.add_run(text)
    set_run_font(run, size=Pt(14 if level == 1 else 12), bold=True)
    return p


def add_body(cell, parts):
    p = cell.add_paragraph()
    set_para_format(p, indent=True)
    for text, citations in parts:
        run = p.add_run(text)
        set_run_font(run)
        for citation in citations:
            cite = p.add_run(f"[{citation}]")
            set_run_font(cite, size=Pt(9), superscript=True)
    return p


def add_numbered_body(cell, number, text):
    p = cell.add_paragraph()
    set_para_format(p, indent=True)
    run = p.add_run(f"{number}. {text}")
    set_run_font(run)
    return p


def set_cell_text(cell, text, size=Pt(12), bold=False, center=True):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def allow_row_break(row):
    tr_pr = row._tr.get_or_add_trPr()
    for tag in ("w:cantSplit", "w:trHeight"):
        for element in list(tr_pr.findall(qn(tag))):
            tr_pr.remove(element)


def remove_pagination_constraints(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    for tag in ("w:keepNext", "w:keepLines", "w:pageBreakBefore"):
        for element in list(p_pr.findall(qn(tag))):
            p_pr.remove(element)


def build():
    actual_hash = hashlib.sha256(REFERENCE.read_bytes()).hexdigest().upper()
    if actual_hash != EXPECTED_HASH:
        raise RuntimeError(f"模板哈希不一致：{actual_hash}")

    doc = Document(REFERENCE)
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

    # 删除模板示例说明，保留学校主标题和表后审核说明。
    if len(doc.paragraphs) > 1:
        p = doc.paragraphs[1]._element
        p.getparent().remove(p)

    title_p = doc.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(10)
    for run in title_p.runs:
        set_run_font(run, chinese="黑体", size=Pt(18), bold=True)

    table = doc.tables[0]
    table.autofit = False
    for row in table.rows:
        allow_row_break(row)

    set_cell_text(table.cell(0, 0), "论文（设计）题目", bold=True)
    set_cell_text(table.cell(0, 1), "基于SpringCloud的票务平台设计与研发", size=Pt(12))
    set_cell_text(table.cell(1, 1), "AI工程师学院")
    set_cell_text(table.cell(1, 3), "软件工程")
    set_cell_text(table.cell(2, 1), "余凯欣")
    set_cell_text(table.cell(2, 3), "202316510149")
    set_cell_text(table.cell(3, 1), "黄子纯")
    set_cell_text(table.cell(3, 2), "职称或学位", size=Pt(10.5), bold=True)
    set_cell_text(table.cell(3, 3), "硕士")
    set_cell_text(table.cell(4, 1), "1.科研课题      2.企业与社会生产实践       3.自选课题", center=False)

    content_cell = table.cell(5, 0)
    clear_cell(content_cell)
    content_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    add_heading(content_cell, "一、选题的依据（选题来源）及意义")
    add_body(content_cell, [
        ("本课题来源于自选课题。演唱会、音乐节等大型营业性演出具有跨城观演比例高、开票流量集中、库存资源有限和交易时效要求高等特点。中国演出行业协会与灯塔专业版发布的行业分析对大型演出的市场趋势、用户流动和消费特征进行了归纳", [1]),
        ("。这些业务特点要求票务平台同时支持活动发布、巡演与场次配置、实名购票、座位及票档库存、抢票排队、订单支付、电子票、退款和入场核验等环节。", []),
    ])
    add_body(content_cell, [
        ("票务业务同时连接C端消费者、主办方运营人员和平台管理人员，任何库存、订单或支付状态偏差都会影响后续出票、退款和核验。围绕实际源码研究微服务拆分、高并发排队、库存一致性、异步消息与前后端协同，可以检验服务边界、数据归属、接口幂等、可靠性和可维护性等软件工程问题，也能为同类高并发交易平台的设计与验证提供参考。", []),
    ])

    add_heading(content_cell, "二、国内外研究综述")
    add_heading(content_cell, "（一）国内研究综述", level=2)
    add_body(content_cell, [
        ("国内近年的相关研究主要围绕微服务拆分、云原生部署、运行监控和并发性能展开。何锋等人设计了一体化性能监控SaaS云服务，用于采集服务状态并识别异常事件", [2]),
        ("；李淑霞等人采用Spring Cloud构建能源互联网营销服务系统，将数据处理与业务能力划分为独立模块", [3]),
        ("；庞长才基于容器编排、API网关和标准化接口研究管理信息系统的微服务化", [4]),
        ("；张健围绕Spring Cloud多层级组件平台分析模块拆分、组件复用与并发处理", [5]),
        ("。上述研究为服务划分、统一入口、可观测性和并发验证提供了直接参考。", []),
    ])
    add_body(content_cell, [
        ("在综合平台和高并发研究方面，张佳佳、谭艳讨论了一体化服务平台的关键技术和实施策略", [6]),
        ("；陈彦成分析了微服务架构下Web应用开发及跨域通信问题", [7]),
        ("；甘日进针对高并发分布式系统研究负载均衡机制", [8]),
        ("。这些成果说明，多模块系统需要同时处理服务调用、流量分配和前后端接口协作。", []),
    ])
    add_body(content_cell, [
        ("在具体行业系统方面，王佳艺梳理了民航旅客服务系统的架构演变", [9]),
        ("；唐俊、邓东杰以气象信息共享平台为对象研究微服务重构与服务治理", [10]),
        ("。现有研究覆盖平台拆分、服务治理和性能优化，但对演出票务中的票档与座位库存、普通抢票与组队抢票、候补分配、支付退款、电子票转赠和入场核验等状态协同仍需结合完整业务链路进行研究。", []),
    ])

    add_heading(content_cell, "（二）国外研究综述", level=2)
    add_body(content_cell, [
        ("国外研究与工程资料更重视云原生服务的实现方式和性能优化。Gutierrez在Spring Boot 3技术著作中系统介绍了Spring Cloud的服务发现、配置、网关与服务调用机制", [11]),
        ("。Guo等人使用Spring Cloud设计智能物流控制平台，将业务能力拆分为可独立部署的服务", [12]),
        ("。Chen等人针对Spring Cloud Alibaba环境下的负载均衡算法开展优化研究", [13]),
        ("。这些成果为票务平台的服务划分、统一网关和高峰流量治理提供了技术依据。", []),
    ])
    add_body(content_cell, [
        ("综合国内外研究可见，微服务可以降低大型平台模块间的直接耦合，但服务拆分后会带来跨服务调用、事务一致性、流量突增和故障恢复问题。本课题将结合现有源码验证这些问题在票务业务中的具体表现，不把框架能力等同于业务完成度。", []),
    ])

    add_heading(content_cell, "三、研究的目标与研究内容")
    add_heading(content_cell, "（一）研究目标", level=2)
    add_body(content_cell, [
        ("本课题拟以Omni票务平台实际源码为研究对象，面向演出票务系统业务链路长、开票流量集中、库存资源有限和跨服务状态协同复杂等问题，从领域建模、微服务边界、功能模块、交易状态控制及并发治理等方面开展设计与研发。课题拟明确B端运营与C端购票的职责边界，建立用户、票务、订单、支付、通知和抢票服务之间可追溯的协作模型；拟完成活动发布、库存锁定、订单创建、支付确认、电子票生成及候补分配等状态约束的设计；拟建立以功能正确性、接口幂等性、库存一致性和并发性能为核心的验证方案。研究过程将以源代码、配置文件、数据库迁移资产、接口调用关系和测试记录为依据，不以框架功能推断系统能力。", []),
    ])
    add_heading(content_cell, "（二）研究内容", level=2)
    add_numbered_body(content_cell, 1, "业务需求与领域模型研究：拟以普通用户、主办方运营人员和平台管理人员为主要角色，分析各角色的用例与权限边界；围绕活动、巡演、场次、场馆、票档、座位、观演人、订单、支付、退款、电子票和候补记录建立领域对象及关联关系，为服务划分和状态设计提供业务依据。")
    add_numbered_body(content_cell, 2, "微服务架构与数据边界研究：拟分析Spring Cloud Gateway统一入口、Nacos服务发现与配置、Sentinel流量治理和OpenFeign内部调用机制；结合prod-split运行拓扑，明确java-user、java-ticket、java-order、java-payment、java-notification及grab-service的职责与数据所有权，设计服务间通过受控接口协作、避免跨服务直接访问数据表的实现方式。")
    add_numbered_body(content_cell, 3, "票务交易状态研究：拟沿活动浏览、实名观演人选择、票档或座位选择、库存锁定、订单快照生成、支付确认、电子票生成、核验、转赠、退款及库存释放链路，分析核心实体状态、状态迁移条件和异常补偿路径，验证订单、支付、票务三方状态能否保持业务一致。")
    add_numbered_body(content_cell, 4, "高并发购票与候补机制研究：拟分析grab-service中的Redis排队、可见库存、请求幂等、自动降档和候补排位逻辑，以及RabbitMQ在候补释放、支付结果和通知处理中的异步协作；研究重复请求、库存竞争、消息重试和服务超时条件下的处理策略。")
    add_numbered_body(content_cell, 5, "服务治理与安全机制研究：拟分析Seata参与跨服务事务协作的边界，结合JWT身份认证、RBAC权限控制、X-Internal-Token内部接口校验、网关限流及审计入口，设计外部访问、服务间调用和后台操作的安全控制闭环。")
    add_numbered_body(content_cell, 6, "前后端业务闭环与系统验证研究：拟核对Next.js前端页面、B端运营入口与后端接口的对应关系，设计功能、边界与并发测试，形成可复现的验证记录。")

    add_heading(content_cell, "（三）拟实现的功能模块", level=2)
    add_numbered_body(content_cell, 1, "C端账户与主体身份模块：拟实现账号认证、注册登录、验证码发送、密码找回与修改、个人资料与头像维护、实名观演人新增修改删除及导出、浏览记录管理、主办方申请与申请进度查询等功能。")
    add_numbered_body(content_cell, 2, "C端内容发现与互动模块：拟实现首页活动展示、活动与巡演详情、分类查询、关键词搜索、历史浏览、订阅日历、活动评价与问答、评价举报、通知列表与通知设置、帮助中心、在线客服会话及转人工等功能。")
    add_numbered_body(content_cell, 3, "C端购票交易与票券模块：拟实现座位与票档查询、购票预约、普通抢票、抢票进度与取消、组队抢票的创建加入确认和策略设置、候补加入排位取消及候补订单处理、订单管理与回收站、支付与退款、电子票展示、入场码生成、票券转赠领取与撤销等功能。")
    add_numbered_body(content_cell, 4, "B端活动与资源配置模块：拟实现运营工作台、活动管理与上下架、活动营销配置和购票人通知、巡演与站次管理、站次配置审核与场馆绑定、场次与票档库存、座位图设计与票档绑定、场馆与座位区域及模板、场馆申请审核、艺人资料提交审核与风险标记、活动搜索索引维护以及公开和私有素材管理等功能。")
    add_numbered_body(content_cell, 5, "B端运营治理模块：拟实现订单查询、退款审核、对账批次与差异处理、验票核销概览与记录、评价问答审核及举报处置、风险事件与风险处置、活动暂停、异常任务认领与关闭、审计日志、主办方管理员与运营人员分配、角色权限管理、客服账号和客服会话管理、会话备注标签快捷回复转接升级及关闭等功能。")
    add_numbered_body(content_cell, 6, "平台支撑模块：拟实现统一网关与路由、JWT身份认证、RBAC授权、X-Internal-Token内部接口校验、Nacos服务注册与配置、Sentinel限流治理、OpenFeign服务调用、Redis缓存排队与库存可见性、RabbitMQ异步消息、Seata事务协作、PostgreSQL分服务数据存储，以及站内通知和按配置启用的短信邮件通道。")

    add_heading(content_cell, "（四）拟解决的关键问题", level=2)
    add_numbered_body(content_cell, 1, "服务边界与数据归属问题：确定各业务服务拥有的数据与对外能力，控制跨服务依赖，避免因共享数据表或越权访问造成模块耦合。")
    add_numbered_body(content_cell, 2, "库存与交易状态一致性问题：研究库存锁定、订单超时、支付确认、退款和库存释放之间的约束关系，重点验证超卖、少卖及状态悬挂等异常场景。")
    add_numbered_body(content_cell, 3, "高并发条件下的幂等与流量治理问题：研究重复提交、并发抢购、消息重复投递和服务超时对系统状态的影响，并通过幂等键、排队、限流、重试及补偿机制控制风险。")
    add_numbered_body(content_cell, 4, "前后端能力一致性问题：以用户可操作的页面和后台入口为验收载体，验证后端能力是否被正确接入，避免仅有接口实现而缺少可完成的业务流程。")

    add_heading(content_cell, "四、研究方法及可行性分析")
    add_heading(content_cell, "（一）研究方法", level=2)
    add_numbered_body(content_cell, 1, "文献研究法：检索并筛选近三年的演出市场、互联网应用、微服务架构和负载均衡研究成果，归纳票务行业的业务特征、微服务平台的常用设计方法及高并发系统的研究重点，为研究问题和评价维度提供理论依据。")
    add_numbered_body(content_cell, 2, "源码证据分析法：以控制器、服务类、配置文件、数据库迁移脚本、消息生产与消费代码、前端路由及测试用例为直接证据，自入口向下追踪活动发布、购票、支付、出票和候补链路，并通过调用关系与数据访问范围核实服务职责。")
    add_numbered_body(content_cell, 3, "领域建模与系统设计法：依据角色用例和业务对象划分领域职责，描述核心实体、状态迁移、服务接口与数据所有权；结合网关、缓存、消息队列和分布式事务的实际接入位置，分析同步调用、异步协作及异常补偿的适用边界。")
    add_numbered_body(content_cell, 4, "实验测试法：采用单元测试、接口测试、微服务边界检查和前端类型检查验证功能与约束；在固定库存和测试数据条件下使用现有JMeter场景对grab-service施加并发负载，记录吞吐量、P95/P99响应时间、错误率、超卖数、重复订单数和最终状态一致性。实验只报告实际测得结果，并保留环境参数和执行日志。")
    add_numbered_body(content_cell, 5, "对照分析法：在不同并发规模、库存数量和异常条件下比较测试结果，结合日志与数据库状态定位性能瓶颈或一致性偏差，再对照源码解释其产生原因及影响范围。")

    add_heading(content_cell, "（二）可行性分析", level=2)
    add_body(content_cell, [
        ("1. 源码基础可行。课题已具备java-gateway、java-user、java-ticket、java-order、java-payment、java-notification和grab-service等前期原型源码，并包含C端购票页面与B端运营入口。账户、活动、巡演场次、场馆座位、抢票候补、订单支付、退款、电子票、客服、风控和审计等拟研究模块均能在控制器、前端路由、配置或测试文件中定位，研究内容具有可核验的实现基础。", []),
    ])
    add_body(content_cell, [
        ("2. 技术条件可行。Java业务服务采用Spring Boot 2.7.18、Spring Cloud 2021.0.8和Spring Cloud Alibaba 2021.0.5.0，抢票服务采用NestJS，前端采用Next.js；项目已接入PostgreSQL、Redis、Nacos、Sentinel、RabbitMQ和Seata。上述组件覆盖数据持久化、缓存、服务治理、异步消息和事务协作等研究环节。", []),
    ])
    add_body(content_cell, [
        ("3. 实验条件可行。项目包含Java与NestJS测试、微服务边界检查脚本、前端类型检查命令以及“1000并发、100库存”的JMeter场景，可用于构建功能、架构边界和并发性能的分层验证。实验阶段将固定软件版本、服务配置、库存初值和测试数据，重复执行关键场景，以降低环境波动对结果的影响。", []),
    ])
    add_body(content_cell, [
        ("4. 风险控制可行。针对本地中间件状态、跨服务调用失败、消息重复投递和测试结果偏差等风险，研究过程将通过启动前环境检查、请求标识与日志追踪、数据库状态核对、幂等与补偿场景测试及多轮结果对照进行控制。若部分功能受外部支付或部署环境限制，则只验证源码能够支撑的本地链路，并在论文中说明验证边界，不推断未执行的实验结果。", []),
    ])

    add_heading(content_cell, "五、进度安排")
    add_body(content_cell, [("论文起止日期和各阶段安排：待学校通知。", [])])

    add_heading(content_cell, "六、主要参考文献")
    for index, ref in enumerate(REFERENCES, 1):
        p = content_cell.add_paragraph()
        fmt = p.paragraph_format
        fmt.left_indent = Pt(24)
        fmt.first_line_indent = Pt(-24)
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.line_spacing = 1.0
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.keep_together = False
        run = p.add_run(f"[{index}] {ref}")
        set_run_font(run, size=Pt(10.5))

    # 删除清空单元格时保留的首个空段落。
    if content_cell.paragraphs and not content_cell.paragraphs[0].text:
        empty = content_cell.paragraphs[0]._element
        empty.getparent().remove(empty)

    for paragraph in content_cell.paragraphs:
        remove_pagination_constraints(paragraph)

    # 审核与签字由教师完成，清除模板中的示例勾选、红字和示例日期。
    signature_cell = table.cell(6, 0)
    clear_cell(signature_cell)
    signature_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    signature_lines = [
        "选题是否合适：  是 □    否 □",
        "课题能否实现：  能 □    不能 □",
        "",
        "指导教师（签名）：________________",
        "开题时间：待学校通知",
    ]
    for index, line in enumerate(signature_lines):
        p = signature_cell.paragraphs[0] if index == 0 else signature_cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index < 3 else WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(line)
        set_run_font(run, size=Pt(12))
    allow_row_break(table.rows[6])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
