from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


SECOND = Path(r"C:\Users\Administrator\Desktop\开题报告\202316510149-余凯欣-开题报告-第2稿.docx")
THIRD = Path(r"C:\Users\Administrator\Desktop\开题报告\202316510149-余凯欣-开题报告-第3稿.docx")


def cell_text(doc):
    return "\n".join(paragraph.text for paragraph in doc.tables[0].cell(5, 0).paragraphs)


second_doc = Document(SECOND)
third_doc = Document(THIRD)
third_text = cell_text(third_doc)
third_tail = third_text.split("三、研究的目标与研究内容", 1)[1]
schedule = third_tail.split("五、进度安排", 1)[1].split("六、主要参考文献", 1)[0]

for row_index in range(5):
    second_row = [cell.text for cell in second_doc.tables[0].rows[row_index].cells]
    third_row = [cell.text for cell in third_doc.tables[0].rows[row_index].cells]
    assert second_row == third_row, f"基础信息第{row_index + 1}行发生变化"

assert schedule.strip() == "论文起止日期和各阶段安排：待学校通知。"
assert "RAG" not in schedule and "大语言模型" not in schedule
assert "（三）拟实现的功能模块" in third_tail
assert "（四）拟解决的关键问题" in third_tail
assert "项目已形成" not in third_tail

required_modules = [
    "账号认证", "实名观演人", "浏览记录", "订阅日历", "活动评价与问答",
    "座位与票档", "预约", "普通抢票", "组队抢票", "候补", "订单管理",
    "支付与退款", "电子票", "票券转赠", "在线客服", "主办方申请",
    "活动管理", "巡演与站次", "场次与票档库存", "场馆与座位", "艺人资料",
    "对账", "验票核销", "评价问答审核", "风险事件", "异常任务", "审计日志",
    "角色权限", "统一网关", "服务注册与配置", "Redis", "RabbitMQ", "Seata",
]
for module in required_modules:
    assert module in third_tail, f"功能模块清单缺少：{module}"

reference_block = third_text.split("六、主要参考文献", 1)[1]
references = re.findall(r"^\[(\d+)] (.+)$", reference_block, flags=re.MULTILINE)
assert [number for number, _ in references] == [str(index) for index in range(1, 14)]
assert sum("[R]" in ref for _, ref in references) == 1, "行业报告数量不是1篇"
assert sum("[J]" in ref for _, ref in references) >= 10, "期刊论文不足10篇"
assert "微服务架构的一体化性能监控SaaS云设计与实现" in reference_block
assert "基于Spring Cloud微服务架构的能源互联网营销服务系统设计" in reference_block
assert "基于云原生技术的管理信息系统微服务架构设计与实现" in reference_block
assert "基于Spring Cloud微服务架构的工业软件多层级组件平台设计" in reference_block
assert not re.search(r"\bdoi\b|https?://doi\.org", third_text, flags=re.IGNORECASE)

body = third_text.split("六、主要参考文献", 1)[0]
for index in range(1, 14):
    assert f"[{index}]" in body, f"正文缺少参考文献[{index}]引用"

content_paragraphs = third_doc.tables[0].cell(5, 0).paragraphs
superscript_citations = set()
reference_started = False
for paragraph in content_paragraphs:
    text = paragraph.text.strip()
    if text == "六、主要参考文献":
        reference_started = True
        continue
    if reference_started:
        assert paragraph.paragraph_format.left_indent == Pt(24)
        assert paragraph.paragraph_format.first_line_indent == Pt(-24)
        continue
    for run in paragraph.runs:
        if run.font.superscript and re.fullmatch(r"\[(\d+)]", run.text):
            superscript_citations.add(int(run.text[1:-1]))
    is_heading = bool(re.match(r"^(一、|二、|三、|四、|五、|六、|（[一二三四五六]）)", text))
    if text and not is_heading:
        assert paragraph.paragraph_format.first_line_indent == Pt(24), f"正文未首行缩进：{text[:30]}"
assert superscript_citations == set(range(1, 14))

for row in third_doc.tables[0].rows:
    tr_pr = row._tr.trPr
    if tr_pr is None:
        continue
    assert not tr_pr.findall(qn("w:cantSplit")), "表格存在禁止跨页断行设置"
    assert not tr_pr.findall(qn("w:trHeight")), "表格存在固定行高设置"

section = third_doc.sections[0]
assert round(section.page_width.cm, 1) == 21.0
assert round(section.page_height.cm, 1) == 29.7

print("结构校验通过：功能模块覆盖源码业务域；正文首行缩进2字符；13篇文献中仅1篇报告、期刊论文不少于10篇；正文引用完整且无DOI；表格允许跨页断行；页面为A4。")
